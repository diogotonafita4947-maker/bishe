import os
import io
import json
import random
import string
import pandas as pd
from bs4 import BeautifulSoup
from urllib.parse import unquote

from django.conf import settings
from django.http import HttpResponse
from django.utils import timezone
from django.contrib.auth.hashers import make_password
from django.db import transaction
from django.db.models import Q 

from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, IsAdminUser
from rest_framework.parsers import MultiPartParser

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING 
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement

from .models import (
    User, ReportTemplate, TrainingTask, StudentReport, ClassGroup, ReportAttachment,
    StudentWhitelist, TeacherWhitelist
)
from .serializers import (
    UserSerializer, ReportTemplateSerializer, TrainingTaskSerializer, 
    StudentReportSerializer, ClassGroupSerializer, ReportAttachmentSerializer,
    UserRegistrationSerializer,
    StudentWhitelistSerializer, TeacherWhitelistSerializer
)

class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer
    permission_classes = [IsAuthenticated]

    @action(detail=False, methods=['get'])
    def me(self, request):
        return Response(self.get_serializer(request.user).data)

    @action(detail=False, methods=['post'], permission_classes=[]) 
    def register(self, request):
        serializer = UserRegistrationSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "注册成功"})
        return Response(serializer.errors, status=400)

    # 1. 获取白名单列表
    @action(detail=False, methods=['get'], url_path='get_whitelist', permission_classes=[IsAdminUser])
    def get_whitelist(self, request):
        role_type = request.query_params.get('type', 'student')
        page = int(request.query_params.get('page', 1))
        page_size = int(request.query_params.get('pageSize', 10))
        keyword = request.query_params.get('keyword', '')

        if role_type == 'student':
            qs = StudentWhitelist.objects.all().order_by('-id')
            if keyword:
                qs = qs.filter(Q(name__icontains=keyword) | Q(student_id__icontains=keyword))
            serializer_class = StudentWhitelistSerializer
        else:
            qs = TeacherWhitelist.objects.all().order_by('-id')
            if keyword:
                qs = qs.filter(Q(name__icontains=keyword) | Q(teacher_id__icontains=keyword))
            serializer_class = TeacherWhitelistSerializer

        total = qs.count()
        start = (page - 1) * page_size
        end = start + page_size
        page_data = qs[start:end]
        
        serializer = serializer_class(page_data, many=True)
        return Response({'total': total, 'list': serializer.data})

    # ★★★ 2. 导入接口 (智能识别版) ★★★
    @action(detail=False, methods=['post'], url_path='upload_roster', parser_classes=[MultiPartParser], permission_classes=[IsAdminUser])
    def upload_roster(self, request):
        file = request.FILES.get('file')
        role_type = request.data.get('type') # 前端可能没传这个，导致为None
        
        if not file: return Response({'error': '请上传Excel文件'}, 400)
        
        try:
            df = pd.read_excel(file).fillna('')
            # 去除表头空格
            df.columns = df.columns.str.strip()
            
            # ★★★ 核心修复：如果前端没传角色，根据表头自动猜！ ★★★
            if not role_type or role_type == 'null' or role_type == 'undefined':
                cols = list(df.columns)
                if '学号' in cols or 'student_id' in cols:
                    role_type = 'student'
                elif '工号' in cols or 'teacher_id' in cols:
                    role_type = 'teacher'
                else:
                    # 默认设为学生，防止报错
                    role_type = 'student'

            print(f">>> 智能识别角色为: {role_type}")
            print(f">>> Excel表头: {list(df.columns)}")
            
            count = 0
            
            def get_val(row, keys):
                for k in keys:
                    if k in row: return str(row[k]).strip()
                return ''

            if role_type == 'student':
                for index, row in df.iterrows():
                    sid = get_val(row, ['学号', 'student_id', 'id'])
                    name = get_val(row, ['姓名', 'name'])
                    college = get_val(row, ['学院', 'college'])
                    major = get_val(row, ['专业', 'major'])
                    
                    if sid and name:
                        StudentWhitelist.objects.update_or_create(
                            student_id=sid,
                            defaults={'name': name, 'college': college, 'major': major}
                        )
                        count += 1
            elif role_type == 'teacher':
                for index, row in df.iterrows():
                    tid = get_val(row, ['工号', 'teacher_id', 'id'])
                    name = get_val(row, ['姓名', 'name'])
                    if tid and name:
                        TeacherWhitelist.objects.update_or_create(
                            teacher_id=tid, defaults={'name': name}
                        )
                        count += 1
            
            print(f">>> 导入完成，共 {count} 条")
            return Response({'message': f'成功导入 {count} 条数据'})
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({'error': f"导入失败: {str(e)}"}, 400)

# ================== 以下内容未变 ==================

class ClassGroupViewSet(viewsets.ModelViewSet):
    queryset = ClassGroup.objects.all(); serializer_class = ClassGroupSerializer; permission_classes = [IsAuthenticated]
    def get_queryset(self):
        user = self.request.user
        if user.is_superuser or user.role == 'admin': return ClassGroup.objects.all()
        if user.role == 'teacher': return ClassGroup.objects.filter(created_by=user)
        return user.classes.all()
    def perform_create(self, serializer):
        code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))
        serializer.save(created_by=self.request.user, invite_code=code)
    @action(detail=False, methods=['post'])
    def join_class(self, request):
        code = request.data.get('invite_code', '').strip().upper()
        group = ClassGroup.objects.filter(invite_code=code).first()
        if not group: return Response({"error": "无效邀请码"}, 400)
        request.user.classes.add(group)
        return Response({"message": "加入成功"})

# ... (上面的代码保持不变)

# ... (保持上面的引用不变)

# ★★★ 修复：真正的 Word 解析逻辑 ★★★
# ... (保留头部引用)

# ... (保持头部引用不变)

class ReportTemplateViewSet(viewsets.ModelViewSet):
    queryset = ReportTemplate.objects.all()
    serializer_class = ReportTemplateSerializer
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser]

    @action(detail=False, methods=['post'])
    def upload_docx(self, request):
        file = request.FILES.get('file')
        if not file: return Response({'error': '请选择文件'}, 400)
        
        try:
            document = Document(file)
            structure = []
            
            for para in document.paragraphs:
                text = para.text.strip()
                if not text: continue
                
                # ★★★ 规则 1：{{ }} 双大括号 -> 老师发布时填写 ★★★
                # 例如 Word 里写：实验目的：{{实验目的}}
                # 系统会生成一个输入框，老师发布任务时填写具体目的
                if '{{' in text and '}}' in text:
                    start = text.find('{{') + 2
                    end = text.find('}}')
                    label = text[start:end] # 提取出 "实验目的"
                    
                    structure.append({
                        "type": "teacher_block", # 标记：老师块
                        "label": label,
                        "value": "" # 默认为空，等老师在网页上填
                    })
                
                # ★★★ 规则 2：【 】中括号 -> 学生做题时填写 ★★★
                elif '【' in text and '】' in text:
                    start = text.find('【') + 1
                    end = text.find('】')
                    label = text[start:end]
                    
                    structure.append({
                        "type": "textarea", # 标记：学生输入框
                        "label": label,
                        "value": ""
                    })
                
                # 规则 3：普通说明文字
                else:
                    structure.append({
                        "type": "paragraph",
                        "label": "说明",
                        "value": text
                    })
            
            template = ReportTemplate.objects.create(
                title=file.name.replace('.docx', '').replace('.doc', ''),
                description="Word 自动导入",
                content_structure=structure
            )
            return Response(ReportTemplateSerializer(template).data)
        except Exception as e:
            import traceback; traceback.print_exc()
            return Response({'error': str(e)}, 500)

# ... (保持下面的代码不变)
class TrainingTaskViewSet(viewsets.ModelViewSet):
    serializer_class = TrainingTaskSerializer; permission_classes = [IsAuthenticated]
    def get_queryset(self): return TrainingTask.objects.all() 
    def perform_create(self, serializer): serializer.save(teacher=self.request.user)
    @action(detail=True, methods=['get'])
    def my_report(self, request, pk=None):
        task = self.get_object()
        report, _ = StudentReport.objects.get_or_create(task=task, student=request.user)
        return Response(StudentReportSerializer(report).data)

class StudentReportViewSet(viewsets.ModelViewSet):
    serializer_class = StudentReportSerializer
    permission_classes = [IsAuthenticated]
    def get_queryset(self):
        user = self.request.user
        if user.role == 'teacher' or user.is_superuser: return StudentReport.objects.all()
        return StudentReport.objects.filter(student=user)
    def perform_create(self, serializer): serializer.save(student=self.request.user)
    def perform_update(self, serializer):
        if self.request.data.get('status') == 'submitted':
            serializer.save(submitted_at=timezone.now(), submit_count=self.get_object().submit_count + 1)
        else: serializer.save()
    @action(detail=True, methods=['post'])
    def grade(self, request, pk=None):
        report = self.get_object()
        report.score = request.data.get('score'); report.feedback = request.data.get('feedback')
        report.status = 'graded'; report.save(); return Response({'status': 'success'})
    @action(detail=True, methods=['post'])
    def reject(self, request, pk=None):
        report = self.get_object(); report.status = 'returned'; report.save(); return Response({'status': 'success'})
    @action(detail=True, methods=['post'], parser_classes=[MultiPartParser])
    def upload_attachment(self, request, pk=None):
        file = request.FILES.get('file')
        if file: ReportAttachment.objects.create(report=self.get_object(), file=file); return Response({'status': 'success'})
        return Response({'error': 'no file'}, 400)
    @action(detail=True, methods=['delete'])
    def remove_attachment(self, request, pk=None):
        ReportAttachment.objects.filter(id=request.data.get('attachment_id')).delete(); return Response({'status': 'success'})

    @action(detail=True, methods=['get'])
    def export_docx(self, request, pk=None):
        try:
            report = self.get_object(); document = Document()
            section = document.sections[0]
            section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5); section.right_margin = Cm(2.0)

            def set_font(run, font_name='黑体', size=12, bold=False):
                run.font.name = font_name
                rPr = run._element.get_or_add_rPr(); rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), font_name); rFonts.set(qn('w:ascii'), font_name); rFonts.set(qn('w:hAnsi'), font_name)
                run.font.size = Pt(size); run.bold = bold
            def set_cell_bg(cell, color_hex="E7E6E6"):
                shading_elm = OxmlElement('w:shd'); shading_elm.set(qn('w:val'), 'clear'); shading_elm.set(qn('w:color'), 'auto'); shading_elm.set(qn('w:fill'), color_hex)
                cell._element.get_or_add_tcPr().append(shading_elm)
            def to_vertical_ancient_style(text):
                if not text: return ""
                replacements = {'(': '︵', ')': '︶', '（': '︵', '）': '︶', '[': '︻', ']': '︼', '【': '︻', '】': '︼', '{': '︷', '}': '︸', '<': '︽', '>': '︾', '《': '︽', '》': '︾'}
                converted_text = text
                for k, v in replacements.items(): converted_text = converted_text.replace(k, v)
                return "\n".join(list(converted_text))
            def set_tight_spacing(paragraph):
                paragraph.paragraph_format.space_before = Pt(0); paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY; paragraph.paragraph_format.line_spacing = Pt(12)
            def insert_image(container, src):
                if not src: return
                try:
                    rel_path = src.split('/media/')[-1] if '/media/' in src else src; rel_path = unquote(rel_path); image_path = os.path.join(settings.MEDIA_ROOT, rel_path)
                    if os.path.exists(image_path):
                        p = container.add_paragraph() if hasattr(container, 'add_paragraph') else container
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER; run = p.add_run(); run.add_picture(image_path, width=Inches(4.2))
                except: pass
            def parse_html_to_cell(cell, html_content):
                if not html_content: return
                try:
                    soup = BeautifulSoup(html_content, 'lxml'); elements = soup.find_all(['p', 'div', 'img'])
                    if not elements and soup.text.strip(): p = cell.add_paragraph(soup.text.strip()); set_font(p.runs[0], '宋体', 11); return
                    for el in elements:
                        if el.name in ['p', 'div']: text = el.get_text().strip(); (p := cell.add_paragraph(text), set_font(p.runs[0], '宋体', 11)) if text else None; [insert_image(cell, img.get('src')) for img in el.find_all('img')]
                        elif el.name == 'img': insert_image(cell, el.get('src'))
                except: cell.add_paragraph(html_content)

            header = section.header; [p._element.getparent().remove(p._element) for p in header.paragraphs]
            header_p = header.add_paragraph(); header_p.alignment = 1; logo_path = os.path.join(settings.MEDIA_ROOT, 'logo.png')
            if os.path.exists(logo_path): header_p.add_run().add_picture(logo_path, height=Inches(0.6))

            college = report.student.college or "________"; major = report.student.major or "________"
            p_sub = document.add_paragraph(); p_sub.alignment = 1
            run_c = p_sub.add_run(f" {college} "); run_c.underline = True; set_font(run_c, '黑体', 12, True)
            run_fix1 = p_sub.add_run(" 学院 "); set_font(run_fix1, '黑体', 12, True)
            run_m = p_sub.add_run(f" {major} "); run_m.underline = True; set_font(run_m, '黑体', 12, True)
            run_fix2 = p_sub.add_run(" 专业学生实"); set_font(run_fix2, '黑体', 12, True)
            p_title = document.add_paragraph(); p_title.alignment = 1; run_title = p_title.add_run("训 (验) 报 告"); set_font(run_title, '黑体', 20, True); p_title.paragraph_format.space_after = Pt(15)

            table = document.add_table(rows=0, cols=6); table.style = 'Table Grid'; table.alignment = 1; table.autofit = False 
            def add_custom_row(height_cm=1.2):
                row = table.add_row(); trPr = row._tr.get_or_add_trPr(); trHeight = OxmlElement('w:trHeight'); trHeight.set(qn('w:val'), str(int(height_cm * 567))); trHeight.set(qn('w:hRule'), "atLeast"); trPr.append(trHeight); return row

            r1 = add_custom_row(); labels1 = ["姓  名", "专业班级", "学  号"]; values1 = [report.student.first_name or report.student.username, report.task.target_class.name, report.student.username]
            for i in range(3):
                lbl = r1.cells[i*2]; val = r1.cells[i*2+1]; set_cell_bg(lbl); lbl.vertical_alignment = 1; val.vertical_alignment = 1
                p = lbl.paragraphs[0]; p.alignment=1; set_font(p.add_run(labels1[i]), '黑体', 12, False)
                p = val.paragraphs[0]; p.alignment=1; set_font(p.add_run(values1[i]), '宋体', 11)

            r2 = add_custom_row(); c0 = r2.cells[0]; set_cell_bg(c0); c0.vertical_alignment=1; p = c0.paragraphs[0]; p.alignment=1; set_font(p.add_run("实训(验)名称"), '黑体', 12, False)
            c_name = r2.cells[1].merge(r2.cells[2]); c_name.vertical_alignment=1; p = c_name.paragraphs[0]; p.alignment=1; set_font(p.add_run(report.task.title), '宋体', 11)
            c3 = r2.cells[3]; set_cell_bg(c3); c3.vertical_alignment=1; p = c3.paragraphs[0]; p.alignment=1; set_font(p.add_run("指导教师"), '黑体', 12, False)
            c_tea = r2.cells[4].merge(r2.cells[5]); c_tea.vertical_alignment=1; t_name = report.task.teacher.first_name or report.task.teacher.username; p = c_tea.paragraphs[0]; p.alignment=1; set_font(p.add_run(t_name), '宋体', 11)

            r3 = add_custom_row(); c_loc_lbl = r3.cells[0]; set_cell_bg(c_loc_lbl); c_loc_lbl.vertical_alignment=1; p = c_loc_lbl.paragraphs[0]; p.alignment=1; set_font(p.add_run("实训(验)地点"), '黑体', 12, False)
            c_loc_val = r3.cells[1].merge(r3.cells[2]); c_loc_val.vertical_alignment=1; loc_text = getattr(report.task, 'location', '计算机实训中心'); p = c_loc_val.paragraphs[0]; p.alignment=1; set_font(p.add_run(loc_text), '宋体', 11)
            c_date_lbl = r3.cells[3]; set_cell_bg(c_date_lbl); c_date_lbl.vertical_alignment=1; p = c_date_lbl.paragraphs[0]; p.alignment=1; set_font(p.add_run("日期时间"), '黑体', 12, False)
            c_date_val = r3.cells[4].merge(r3.cells[5]); c_date_val.vertical_alignment=1; d_str = report.submitted_at.strftime("%Y-%m-%d") if report.submitted_at else "未提交"; p = c_date_val.paragraphs[0]; p.alignment=1; set_font(p.add_run(d_str), '宋体', 11)

            raw_struct = report.task.template.content_structure; saved_content = report.content_data; teacher_filled = report.task.task_context or {}
            for item in raw_struct:
                item_type = item.get('type')
                if item_type not in ['teacher_block', 'textarea']: continue 
                label = item.get('label', ''); is_vertical = ("步骤" in label or "过程" in label or "结论" in label or "心得" in label); row = add_custom_row(height_cm=8.0 if is_vertical else 2.5)
                lbl_c = row.cells[0]; set_cell_bg(lbl_c); lbl_c.vertical_alignment=1; p = lbl_c.paragraphs[0]; p.alignment=1
                if is_vertical: run = p.add_run(to_vertical_ancient_style(label)); set_font(run, '黑体', 12, False); set_tight_spacing(p)
                else: run = p.add_run(label); set_font(run, '黑体', 12, False)
                val_c = row.cells[1].merge(row.cells[5]); val_c._element.clear_content(); val_c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP 
                if item_type == 'teacher_block': val = teacher_filled.get(label, item.get('value', '')); p = val_c.add_paragraph(val); set_font(p.runs[0], '宋体', 11)
                elif item_type == 'textarea':
                    val = ""; [val := s_item.get('value', '') for s_item in saved_content if s_item.get('label') == label]
                    if val: parse_html_to_cell(val_c, val)
                    else: p = val_c.add_paragraph("(此处未填写)"); set_font(p.runs[0], '宋体', 11)

            if report.status == 'graded':
                r_comment = add_custom_row(height_cm=4.0); lbl_c = r_comment.cells[0]; set_cell_bg(lbl_c); lbl_c.vertical_alignment=1; p = lbl_c.paragraphs[0]; p.alignment=1; run = p.add_run(to_vertical_ancient_style("教师评语")); set_font(run, '黑体', 12, False); set_tight_spacing(p)
                val_c = r_comment.cells[1].merge(r_comment.cells[5]); val_c._element.clear_content(); val_c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP; p_feed = val_c.add_paragraph(report.feedback or "无"); set_font(p_feed.runs[0], '宋体', 11)
                r_bottom = add_custom_row(height_cm=1.2)
                c0 = r_bottom.cells[0]; set_cell_bg(c0); c0.vertical_alignment=1; p = c0.paragraphs[0]; p.alignment=1; set_font(p.add_run("评  分"), '黑体', 12, False)
                c1 = r_bottom.cells[1]; c1.vertical_alignment=1; p = c1.paragraphs[0]; p.alignment=1; (run_s := p.add_run(str(report.score)), set_font(run_s, 'Arial', 12, True)) if report.score is not None else None
                c2 = r_bottom.cells[2]; set_cell_bg(c2); c2.vertical_alignment=1; p = c2.paragraphs[0]; p.alignment=1; set_font(p.add_run("学  生"), '黑体', 12, False)
                c3 = r_bottom.cells[3]; c3.vertical_alignment=1
                c4 = r_bottom.cells[4]; set_cell_bg(c4); c4.vertical_alignment=1; p = c4.paragraphs[0]; p.alignment=1; set_font(p.add_run("指导教师"), '黑体', 12, False)
                c5 = r_bottom.cells[5]; c5.vertical_alignment=1

            f = io.BytesIO(); document.save(f); f.seek(0)
            from django.utils.encoding import escape_uri_path; filename = f"{report.task.title}-{report.student.first_name}.docx"; response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'); response['Content-Disposition'] = f'attachment; filename="{escape_uri_path(filename)}"'; return response
        except Exception as e: import traceback; traceback.print_exc(); return Response({'error': str(e)}, status=500)